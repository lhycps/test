# encoding: GBK
import datetime
import re
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.shared import Cm, Pt, Inches
import difflib
from docx.shared import Pt
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from openpyxl import load_workbook
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_UNDERLINE
from collections import defaultdict
import os
from djcp.utils import insert_Scheme_confirmation_page
from crm import settings


def add_column_at01(table, new_column_index, column_with):
    """
��ָ���в���հ���
    @param table: doc�б��
    @param new_column_index:�ڵڶ��мӣ���ʼֵ��0��������ڵڶ��в壬��д1
    @param column_with:�����пCm(2)
    @return:
    """
    if new_column_index >= len(table.columns):
        return table.add_column(column_with)

    tbl_grid = table._tbl.tblGrid
    grid_col = tbl_grid.add_gridCol()
    grid_col.w = column_with
    tbl_grid.insert(new_column_index, grid_col)
    for tr in table._tbl.tr_lst:
        tc = tr.add_tc()
        tc.width = column_with
        real_tcs = []
        for tc_index, tc_obj in enumerate(tr):
            if tc_obj.__class__.__name__ == "CT_Tc":
                real_tcs.append({
                    "tc": tc_obj,
                    "index": tc_index,
                })
        tr.insert(real_tcs[new_column_index]["index"], tc)
    return table.columns[new_column_index]


def add_column_at(table, new_column_index, column_with):
    """
��ָ���в���հ���
    @param table: doc�б��
    @param new_column_index:�ڵڶ��мӣ���ʼֵ��0��������ڵڶ��в壬��д1
    @param column_with:�����пCm(2)
    @return:
    """
    if new_column_index >= len(table.columns):
        return table.add_column(column_with)
    for tr in table._tbl.tr_lst:
        real_tcs = []
        for tc_index, tc_obj in enumerate(tr):
            real_tcs.append({
                "tc": tc_obj,
                "index": tc_index,
            })
        tc = tr.add_tc()
        tc.width = column_with
        tr.insert(new_column_index + 1, tc)
    return table.columns[new_column_index + 1]


def insert_objecte(doc, sheetnum):
    """
����һ�е�word���ڱ�ͷд��"����������"
    :param doc:
    """
    table = doc.tables[sheetnum]
    # col = table.columns[3]  # ��λ��
    # col_cells = table.add_column(Cm(2)).cells  # ��word������һ�м���һ��
    col_cells = add_column_at01(table, 1, Cm(2)).cells
    cell = col_cells[0]
    cell.text = '����������'
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10.5)
    # bg = parse_xml(r'<w:shd{nsdecls("w")} w:fill="888888"/>')
    bg = parse_xml(r'<w:shd {} w:fill="{color_value}"/>'.format(nsdecls('w'), color_value='#A6A6A6'))
    cell._tc.get_or_add_tcPr().append(bg)
    return table


def insert_number(doc, C, index_wordlist, excel_list2, index_excellist, cellmax, sheetnum):
    '''�����ŵķ���'''
    table = insert_objecte(doc, sheetnum)
    cols = len(table.columns)
    dd = defaultdict(list)
    for k, va in [(v, i) for i, v in enumerate(index_excellist)]:
        dd[k].append(va)
    dd = dict(dd)
    for k, v in dd.items():
        # print('k:{},v:{}'.format(k,v))
        if len(v) == 1:
            run = table.cell(index_wordlist[v[0]], 1).paragraphs[0].add_run(
                'LISS-' + excel_list2[index_excellist[v[0]]] + '-001')
            run.font.name = '���ķ���'
            run.font.size = Pt(10.5)

        else:
            for index, value in enumerate(v):
                if index < 9:
                    run = table.cell(index_wordlist[value], 1).paragraphs[0].add_run(
                        'LISS-' + excel_list2[index_excellist[value]] + '-00' + str(index + 1))
                    run.font.name = '���ķ���'
                    run.font.size = Pt(10.5)
                elif index < 99:
                    run = table.cell(index_wordlist[value], 1).paragraphs[0].add_run(
                        'LISS-' + excel_list2[index_excellist[value]] + '-0' + str(index + 1))
                    run.font.name = '���ķ���'
                    run.font.size = Pt(10.5)
                else:
                    run = table.cell(index_wordlist[value], 1).paragraphs[0].add_run(
                        'LISS-' + excel_list2[index_excellist[value]] + '-' + str(index + 1))
                    run.font.name = '���ķ���'
                    run.font.size = Pt(10.5)

    for k, v in enumerate(C):
        if k < 9:
            run = table.cell(v, 1).paragraphs[0].add_run('LISS-' + cellmax + '-00' + str(k + 1))
            run.font.name = '���ķ���'
            run.font.size = Pt(10.5)
        elif k < 99:
            run = table.cell(v, 1).paragraphs[0].add_run('LISS-' + cellmax + '-0' + str(k + 1))
            run.font.name = '���ķ���'
            run.font.size = Pt(10.5)
        else:
            run = table.cell(v, 1).paragraphs[0].add_run('LISS-' + cellmax + '-' + str(k + 1))
            run.font.name = '���ķ���'
            run.font.size = Pt(10.5)


def find_sameorDiff_Number(index_list, index_wordlist):
    A = set(index_list).intersection(set(index_wordlist))  # ����
    B = set(index_list).union(set(index_wordlist))  # ����
    C = set(index_list).difference(set(index_wordlist))  # �����list1�е�����list2�е�Ԫ��
    D = set(index_wordlist).difference(set(index_list))  # �����list2�е�����list1�е�Ԫ��
    # print("����Ԫ�ظ���:" + str(len(A)))
    # print("����Ԫ�ظ���:" + str(len(B)))
    # print("��list1�е�����list2�е�Ԫ�ظ���:", C)
    # print("��list2�е�����list1�е�Ԫ�ظ���:" + str(len(D)))
    C = list(C)
    return A, B, C, D


def mate(excel_list1, word_list2):
    index_wordlist = []
    excelnew_list = []
    index_excellist = []
    index_list = []
    for index, item in enumerate(word_list2):
        index = index + 1
        index_list.append(index)
        for index_excel, j in enumerate(excel_list1):

            seq = difflib.SequenceMatcher(None, item, j)
            ratio = seq.ratio()
            if ratio > 0.2:
                index_wordlist.append(index)
                excelnew_list.append(j)
                index_excellist.append(index_excel)
                # print('excel������', index_excel)

    return index_wordlist, excelnew_list, index_excellist, index_list


def read_word(doc, sheetnum):
    """
��ȡword�е�����
    :param doc:
    :return:word���ڶ��е�����
    """
    table = doc.tables[sheetnum]
    rows = len(table.rows)  # ��ȡ�������
    word_list2 = []
    for row in range(1, rows):
        col = table.cell(row, 1).text
        word_list2.append(col)
    return word_list2


def read_excel(wb):
    """
��ȡexcel�е�����
    :param wb: excel������
    :return: excel����һ�е������б�
    """
    ws = wb.active
    rowmax = ws.max_row
    colmax = ws.max_column
    cellmax = ws.cell(rowmax, colmax).value
    excel_list1 = []
    excel_list2 = []
    for row in range(1, ws.max_row + 1):
        value1 = ws.cell(row, 1).value
        value2 = ws.cell(row, 2).value
        excel_list1.append(value1)
        excel_list2.append(value2)

    return excel_list1, excel_list2, cellmax


def excel_col2data(wb, index_excellist):
    ws = wb.active
    for index_excel in index_excellist:
        print(index_excel)


def management_system():
    """
��������ȫ�����Ա����ȫ�����ƶȵ�ֻ��һ����ŵĴ�����
    """
    word_list2 = read_word(doc, sheetnum)
    excel_list1, excel_list2, cellmax = read_excel(wb)
    t = len(word_list2)
    table = insert_objecte(doc, sheetnum)
    cols = len(table.columns)
    for k, v in enumerate(word_list2):
        if k < 9:
            run = table.cell(k + 1, 1).paragraphs[0].add_run('LISS-' + cellmax + '-00' + str(k + 1))
            run.font.name = '���ķ���'
            run.font.size = Pt(10.5)
        else:
            run = table.cell(k + 1, 1).paragraphs[0].add_run('LISS-' + cellmax + '-0' + str(k + 1))
            run.font.name = '���ķ���'
            run.font.size = Pt(10.5)


def network_equipment():
    """
�����豸�Ĵ���������Ϊ�����ϵ�ϵͳ�������豸�����������豸��������Ҫ������������
    """
    word_list2 = read_word(doc, sheetnum)
    if '�����治�漰' in word_list2:
        table = insert_objecte(doc, sheetnum)
        cols = len(table.columns)
        table.cell(1, 0).merge(table.cell(1, cols - 1))
    elif len(word_list2) == 0:
        table = insert_objecte(doc, sheetnum)
        cols = len(table.columns)
        for i in range(1):  # ģ�����Ѿ���һ���ˣ������ܹ�ֻ������len(supplier)��
            table.add_row()
        table.cell(1, 0).merge(table.cell(1, cols - 1))  # �ϲ���Ԫ��
        run = table.cell(1, 0).paragraphs[0].add_run('��ϵͳ���漰')
        run.font.name = '���ķ���'
        run.font.size = Pt(10.5)
        table.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    else:
        start_up()


def insert_row(doc):
    """
������ֻ��һ�еı��������һ�м��ϡ���ϵͳ���漰�����ж���
    :param doc: ���������û�����
    """
    tbs = doc.tables
    for tb in tbs:
        if len(tb.rows) == 1:
            tb.add_row()
            cols = len(tb.columns)
            tb.cell(1, 0).merge(tb.cell(1, cols - 1))  # �ϲ���Ԫ��
            run = tb.cell(1, 0).paragraphs[0].add_run('��ϵͳ���漰')
            run.font.name = '���ķ���'
            run.font.size = Pt(10.5)
            tb.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif len(tb.rows) == 2 and tb.cell(1, 0).text == '':
            cols = len(tb.columns)
            tb.cell(1, 0).merge(tb.cell(1, cols - 1))  # �ϲ���Ԫ��
            run = tb.cell(1, 0).paragraphs[0].add_run('��ϵͳ���漰')
            run.font.name = '���ķ���'
            run.font.size = Pt(10.5)
            tb.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def replace_para(doc, old_text, newtext):
    """
���������ġ��滻�ɡ�����������
    :param doc: ���������û�����
    :param old_text: ��������
    :param newtext: ��������
    """
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:

                if old_text in i.text:
                    text = i.text.replace(str(old_text), str(newtext))
                    i.text = text


def start_up():
    word_list2 = read_word(doc, sheetnum)
    excel_list1, excel_list2, cellmax = read_excel(wb)
    index_wordlist, excelnew_list, index_excellist, index_list = mate(excel_list1, word_list2)
    A, B, C, D = find_sameorDiff_Number(index_list, index_wordlist)
    insert_number(doc, C, index_wordlist, excel_list2, index_excellist, cellmax, sheetnum)


def get_parindex(doc):
    """
Ϊ�˶�λ�������������ľ���λ�ã��ҵ���Ӧ��paragraph.
    :param doc:
    :return:
    """
    paragraphs = list(doc.paragraphs)
    # print(len(paragraphs))
    for i in range(len(paragraphs)):
        paragraph = paragraphs[i]
        if '��������' == paragraph.text:
            return i


def getpara_wordcontent(i):
    """
�������������ݲ��뵽word��(�������)
    """

    paragraphs = list(doc.paragraphs)
    paragraphslist = get_doc01_con()

    for par in doc.paragraphs:
        if '��������' == par.text:
            par = paragraphs[i + 1]
            for parh in paragraphslist:
                par.insert_paragraph_before(parh)


def get_doc01_con():
    """
��ȡdoc01�������������ĵ������������������
    :return:
    """
    paragraphslist = []
    for para in doc01.paragraphs:
        p = para.text
        paragraphslist.append(p)
    return paragraphslist


def getpara_wordcontent02():
    """
�������������ݲ��뵽word��(�������)
    """

    '''��λ����'''
    restr_start = '�ȼ�����������������ͼ'
    pos_start = get_position(restr_start)
    # doc.paragraphs[pos_start - 2].clear()  # ɾ��ָ������

    standard_str_start = doc.paragraphs[pos_start + 1].text  # �������ݵĶ�λλ��
    restr_end = '����ȼ������������'
    pos_end = get_position(restr_end)
    standard_str_end = doc.paragraphs[pos_end].text  # �������ݵĶ�λλ��
    print('��ʼ��{},���ݣ�{}��������{},���ݣ�{}'.format(pos_start + 1, standard_str_start, pos_end, standard_str_end))

    '''ɾ������'''
    parlist = []  # ɾ��ָ������
    for i in range(pos_start + 2, pos_end):
        paragraph = doc.paragraphs[i]
        print('��Ҫɾ��������Ϊ', paragraph.text)
        parlist.append(paragraph)

    for j in range(len(parlist)):
        delete_paragraph(parlist[j])

    '''��������'''
    paragraphslist = []
    for para in standard_doc.paragraphs:
        p = para.text
        paragraphslist.append(p)

    paragraphs = list(doc.paragraphs)

    for parh in paragraphslist:  # ��������
        paragraphs[pos_start + 2].insert_paragraph_before(parh)


def delete_paragraph(paragraph):
    '''ɾ��ָ������'''
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def get_standard_con():
    """
��ȡstandard_doc���������ݣ��ĵ������������������
    :return:
    """
    paragraphslist = []
    for para in standard_doc.paragraphs:
        p = para.text
        paragraphslist.append(p)
    return paragraphslist


def get_position(restr):
    '''
    ͨ���ַ������ж�λ
    @param restr:
    @return:
    '''
    pos = 0
    for i in range(len(doc.paragraphs)):
        matchRet = re.findall(restr, doc.paragraphs[i].text)
        if len(matchRet) > 0:
            pos = i
    return pos


def getday(y=2017, m=8, d=15, n=0):
    '''
    ��ȡָ�����ڵ�ǰ����ͺ�������
    @param y: ��
    @param m: ��
    @param d: ��
    @param n: +5����5�죬-5ǰ5��
    @return:
    '''
    the_date = datetime.datetime(y, m, d)
    result_date = the_date + datetime.timedelta(days=n)
    d = result_date.strftime('%Y-%m-%d')
    return d


def time_content():
    '''
    ��ò���������ʱ����뵽������
    ����ͱ�һ����
    @return:
    '''
    '''��λ���ҳ�������������λ��'''
    year_location = '\d{4}[\.\/��-]{,3}'
    month_location = "\d{1,2}[\.\/��-]{,3}"
    day_location = "\d{1,2}[\.\/��-]{,3}"
    casestr = ''
    for i in range(len(doc.paragraphs)):
        mm = re.findall(r"{}��{}��{}�ա�{}��{}�գ��������ƹ��̡�".format(year_location, month_location, day_location, month_location,
                                                           day_location, ), doc.paragraphs[i].text)

        if len(mm) > 0:
            casestr = mm[0]

    case_year = casestr[0:4]  # �õ�������������
    case_mon = casestr[12:14]  # �õ������������º���
    case_day = casestr[15:17]  # �õ�������������

    '''����ʹ��'''
    case_time_face = case_year + '��' + case_mon + '��' + case_day + '��' + ' ' * 30  # ��������
    print(case_time_face)
    tb = doc.tables[0]
    para = tb.cell(3, 1).paragraphs[0]
    para.text = para.text.replace(para.text, case_time_face)
    para.runs[0].bold = True
    para.runs[0].font.size = Pt(16)
    para.runs[0].underline = WD_UNDERLINE.THICK

    '''��һʹ��'''
    case_year = int(case_year)
    case_mon = int(case_mon)
    case_day = int(case_day)
    organization_time = getday(case_year, case_mon, case_day, -1)
    auditd_time = getday(case_year, case_mon, case_day, 0)
    doc.tables[1].rows[5].cells[0].tables[0].rows[6].cells[4].text = organization_time  # ��������
    doc.tables[1].rows[5].cells[0].tables[0].rows[6].cells[4].paragraphs[0].runs[0].font.size = Pt(10.5)
    doc.tables[1].rows[5].cells[0].tables[0].rows[6].cells[4].paragraphs[
        0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.tables[1].rows[5].cells[0].tables[0].rows[7].cells[4].text = auditd_time  # �������
    doc.tables[1].rows[5].cells[0].tables[0].rows[7].cells[4].paragraphs[0].runs[0].font.size = Pt(10.5)
    doc.tables[1].rows[5].cells[0].tables[0].rows[7].cells[4].paragraphs[
        0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.tables[1].rows[5].cells[0].tables[0].rows[8].cells[4].text = auditd_time  # ��׼����
    doc.tables[1].rows[5].cells[0].tables[0].rows[8].cells[4].paragraphs[0].runs[0].font.size = Pt(10.5)
    doc.tables[1].rows[5].cells[0].tables[0].rows[8].cells[4].paragraphs[
        0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def finalstart_up(paths_case):
    """
���մ������
    :param paths_case: �û���������ԭʼ�Ĳ�������
    """
    global doc, sheetnum, wb, doc01, table, standard_doc

    paths_excel = settings.ASSET_NUM
    path_word01 = settings.EVA_METHOD  # ��������·��
    path_word02 = settings.CASE_SIGN  # ����ǩ��·��
    standard_path = settings.STANDARD_PATH  # ��������·��
    doc01 = Document(path_word01)  # ����������doc
    standard_doc = Document(standard_path)  # �������ݵ�doc
    doc_target = Document(path_word02)
    # paragraph_target = doc_target.paragraphs[0]
    table = doc_target.tables[0]
    path = paths_case
    doc = Document(path)
    i = get_parindex(doc)
    getpara_wordcontent(i)  # �������������ݲ��뵽word��(�������)
    getpara_wordcontent02()  # ������������Ҫ�ο��ı�׼
    for name in os.listdir(paths_excel):
        if name == '01����.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 3
            management_system()

        elif name == '02�����豸.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 4
            network_equipment()
        elif name == '03��ȫ�豸.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 5
            start_up()
        elif name == '04�������洢�豸.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 6
            start_up()
        elif name == '05�ն��豸.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 7
            start_up()
        elif name == '06ϵͳ�������ƽ̨.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 8
            start_up()
        elif name == '07ҵ��Ӧ�����ƽ̨.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 9
            start_up()
        elif name == '08��ȫ�����Ա.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 11
            management_system()
        elif name == '09��ȫ�����ĵ�.xlsx':
            path_excel = paths_excel + '/' + name
            wb = load_workbook(path_excel)
            sheetnum = 12
            management_system()
    insert_row(doc)
    caselist = {'6.': '6.0', '��������': '��������', '���簲ȫ�ȼ���������': '���簲ȫ�ȼ��������������ͼƻ�'}
    for k, v in caselist.items():
        replace_para(doc, k, v)
    # insert_Scheme_confirmation_page.getsheet_wordcontent(doc)
    table01 = doc.tables[0]
    a = table01.cell(1, 1).text
    '''ɾ��ָ�����������'''
    insert_Scheme_confirmation_page.del_parcontent(doc)  # ���У�2022��07��12���ٿ�����Ŀ�������飬ȷ���˹�����������Ŀ��Ա������2022��07��23���ٿ�����
    '''�����������ǩ��ȷ��ҳ��'''
    insert_Scheme_confirmation_page.insert_table_after_text(doc, table)
    table.cell(0, 1).text = a
    for r, row in enumerate(table.rows):  # ����������
        for c, cell in enumerate(row.cells):  # ����ÿһ�е���
            cell1 = (r, c)  # �ڵ�Ԫ����д�뵱ǰ��λ�ã�����С��У�
            if cell1 == (0, 1):
                p = cell.paragraphs[0]
                p.runs[0].bold = False
                p.runs[0].font.size = Pt(14)
    try:
        time_content()
    except:
        print('�쳣')

    temp = path.split('/')
    temp1 = temp[-1]
    temp2 = temp1.split('_')

    output_filename = temp2[0] + '_' + temp2[1] + '_' + temp2[2] + '�ͼƻ�'
    output_case_path = os.path.join(settings.CASE_OUTPUT_PATH, '{}.docx'.format(output_filename))
    doc.save(output_case_path)
    print('{}.docx������ɡ�'.format(output_filename))
    return output_filename


if __name__ == '__main__':
    path = 'caseUpload/LISS2020DJCP060_���弯�ŵ���ƽ̨_��������_20220901134536258����ࣩ.docx'
    path1 = 'caseUpload/LISS2022DJCP031_��ƵǼǹ���ƽ̨_��������_20220830124455139���_rLNiWOK.docx'
    path3 = 'caseUpload/LISS2022DJCP068_������֤֤���쳡������ƽ̨_��������_20220901134454287����ࣩ.docx'
    path4 = 'caseUpload/LISS2022DJCP060_������������վȺ_��������_20220901134615994����ࣩ.docx'
    paths_case = os.path.join(settings.MEDIA_ROOT, path4)
    paths_case = re.sub('\\\\', '/', paths_case)  # ·������ת��
    finalstart_up(paths_case)

    # get_case_time(paths_case)
