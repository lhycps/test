from crm import settings
from openpyxl import load_workbook
from docx import Document
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import shared


def sheet01_docx():
    """
表二的内容：安全相关人员
    :param doc: 信息调查表模板
    :param bumen: 部门名称
    """

    # wb = load_workbook(path_docx)
    for sheet_name in wb.sheetnames:
        if '安全相关人员' in sheet_name:
            ws = wb[sheet_name]
            data = []
            for row in range(2, ws.max_row + 1):
                serial_excel = ws["A" + str(row)].value  # 序号
                name_excel = ws["B" + str(row)].value  # 姓名
                post_excel = ws["C" + str(row)].value  # 岗位角色
                number_excel = ws["D" + str(row)].value  # 联系方式
                unit_excel = ws["E" + str(row)].value  # 所属单位
                info = [serial_excel, name_excel, post_excel, number_excel, unit_excel]
                data.append(info)

            table = doc.tables[1]  # 已确定是第二个表格，其索引是1
            serial = []  # 存储序号
            name = []  # 存储人名
            post = []  # 存储职位
            number = []  # 存储电话号码
            unit = []  # 存储所属单位
            for i in data:
                serial.append(i[0])
                name.append(i[1])
                post.append(i[2])
                number.append(i[3])
                unit.append(i[4])

            for i in range(len(name)):  # 模板中已经有一行了，所以总共只需增加len(supplier)行
                table.add_row()
            for row in range(1, len(name) + 1):
                for col in range(6):
                    if col == 0:
                        table.cell(row, 0).text = str(data[row - 1][0])
                    elif col == 1:
                        table.cell(row, 1).text = str(data[row - 1][1])
                    elif col == 2:
                        table.cell(row, 2).text = str(data[row - 1][2])
                    elif col == 3:
                        table.cell(row, 3).text = str(data[row - 1][3])
                    else:
                        table.cell(row, 4).text = str(data[row - 1][4])

    wb.close()

    print('表2:安全相关人员表格处理【完成】')


def sheet02_docx():
    """
信息调查表表三的内容：物理环境情况
    :param doc: 信息调查表模板
    :param systemname: 系统名称
    :return: 机房物理位置
    """
    global loc_excel
    try:
        # wb = load_workbook(path_docx)
        for sheet_name in wb.sheetnames:
            if '机房' in sheet_name:
                ws = wb[sheet_name]
                data = []
                for row in range(2, ws.max_row + 1):
                    serial_excel = ws["A" + str(row)].value  # 序号
                    IDC_excel = ws["B" + str(row)].value  # 物理环境名称
                    location_excel = ws["C" + str(row)].value  # 物理位置
                    appname_excel = appname  # 涉及的信息系统
                    is_important = str(ws["D" + str(row)].value)[:-3]
                    info = [serial_excel, IDC_excel, location_excel, appname_excel, is_important]
                    data.append(info)
                loc_excel = data[0][2]
                print(loc_excel)
                max_row = ws.max_row
                table = doc.tables[3]
                for i in range(max_row - 1):  # 模板中已经有一行了，所以总共只需增加len(supplier)行
                    table.add_row()
                for row in range(1, max_row):
                    for col in range(5):
                        table.cell(row, col).text = str(data[row - 1][col])
        wb.close()
        print('表4：物理机房处理【完成】')
        return loc_excel

    except:
        print('-----异常----')
        return loc_excel, '物理位置为空'


def sheet04_docx(bumen):
    """
信息调查表表五的内容：信息系统网络结构（环境）情况
    :param doc: 信息调查表模板
    :param systemname: 系统名称
    :return: None
    """
    table = doc.tables[4]

    for row in range(1, 4):
        for col in range(2, 11):
            if col == 2:  # 主要业务和信息描述
                table.cell(row, 2).text = appname
            elif col == 3:
                table.cell(row, 3).text = '/'  # IP网段地址
            elif col == 4:
                table.cell(row, 4).text = str(fuwuqinum)  # 服务器数量
            elif col == 5:
                table.cell(row, 5).text = str(zhduangnum)  # 终端数量
            elif col == 9:
                table.cell(row, 9).text = bumen  # 责任部门
            elif col == 10:
                table.cell(row, 10).text = '/'  # 备注
    print('表5：信息系统承载业务（服务）情况处理【完成】')


def sheet05_docx():
    """
表7：外联线路及设备端口（网络外联情况）情况
    :param doc: 信息调查表模板
    :param location_excel: 机房物理位置
    """

    # wb = load_workbook(path_docx)
    for sheet_name in wb.sheetnames:
        if '区域边界' in sheet_name:
            ws = wb[sheet_name]
            data = []
            for row in range(2, ws.max_row + 1):
                serial_excel = ws["A" + str(row)].value  # 序号
                name_excel = ws["B" + str(row)].value  # 边界名称
                sysversion_excel = ws["C" + str(row)].value  # 连接对象名称
                moder_excel = ws["D" + str(row)].value  # 承载的业务应用
                info = [serial_excel, name_excel, sysversion_excel, moder_excel]
                data.append(info)
            max_row = ws.max_row
            table = doc.tables[5]
            for i in range(max_row - 1):  # 模板中已经有一行了，所以总共只需增加len(supplier)行
                table.add_row()
            for row in range(1, max_row):
                for col in range(9):
                    if col == 0:
                        table.cell(row, 0).text = str(data[row - 1][0])
                    elif col == 1:
                        table.cell(row, 1).text = str(data[row - 1][1])
                    elif col == 2:
                        table.cell(row, 2).text = '服务器区'
                    elif col == 3:
                        table.cell(row, 3).text = '/'
                    elif col == 4:
                        if data[row - 1][2] == None:
                            table.cell(row, 4).text = '/'
                        else:
                            table.cell(row, 4).text = str(data[row - 1][2])
                    elif col == 5:
                        table.cell(row, 5).text = '200M'
                    elif col == 6:
                        table.cell(row, 6).text = '防火墙'
                    elif col == 7:
                        table.cell(row, 7).text = appname
                    elif col == 8:
                        table.cell(row, 8).text = '/'
    wb.close()
    print('表6：外联线路及设备端口（网络边界）情况处理【完成】')


def sheet06_docx():
    """
表8：网络互联设备情况
    :param doc: 信息调查表模板
    :param location_excel: 机房物理位置
    """
    # wb = load_workbook(path_docx)
    wangluoquyu = '外联区'
    for sheet_name in wb.sheetnames:
        if '网络设备' in sheet_name:
            ws = wb[sheet_name]
            if ws.max_row > 1:
                data = []
                for row in range(2, ws.max_row + 1):
                    serial_excel = ws["A" + str(row)].value  # 序号
                    name_excel = ws["B" + str(row)].value  # 网络设备名称
                    is_virtual = ws["C" + str(row)].value  # 是否是虚拟设备
                    if is_virtual == 'O':
                        is_virtual = '否'
                    else:
                        is_virtual = '是'
                    sysversion_excel = ws["E" + str(row)].value  # 型号
                    moder_excel = ws["G" + str(row)].value  # ip地址、掩码、网关
                    if moder_excel == '':
                        moder_excel = '/'
                    soft_excel = ws["D" + str(row)].value  # 系统软件级版本
                    number_excel = ws["H" + str(row)].value  # 数量
                    user_excel = ws["F" + str(row)].value  # 主要用途
                    is_hot = '否'  # 是否热备
                    if int(number_excel) > 1:
                        is_hot = '是'
                    impot_excel = str(ws["I" + str(row)].value)[:-3]  # 重要程度
                    mask = '/'
                    info = [serial_excel, name_excel, is_virtual, sysversion_excel, loc_excel, wangluoquyu, moder_excel,
                            soft_excel, number_excel, user_excel, is_hot, impot_excel, mask]
                    data.append(info)
                max_row = ws.max_row
                # print(max_row)
                table = doc.tables[6]
                for i in range(max_row - 1):  # 模板中已经有一行了，所以总共只需增加len(supplier)行
                    table.add_row()
                for row in range(1, max_row):
                    for col in range(13):
                        table.cell(row, col).text = str(data[row - 1][col])
                wb.close()
                break
            else:
                table = doc.tables[6]
                for i in range(1):  # 模板中已经有一行了，所以总共只需增加len(supplier)行
                    table.add_row()
                table.cell(1, 0).merge(table.cell(1, 11))  # 合并单元格
                table.cell(1, 0).paragraphs[0].add_run('本系统不涉及')
                table.cell(1, 0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 居中对齐
    wb.close()
    print('表7：网络互联设备情况情况处理【完成】')


def sheet07_docx():
    """
表9：安全设备的内容
    :param doc: 信息调查表模板
    :param location_excel: 机房物理位置
    """
    sushuwnagluoquyu_excel = '外联区'  # 所属网络区域
    # wb = load_workbook(path_docx)
    for sheet_name in wb.sheetnames:
        if '安全设备' in sheet_name:
            ws = wb[sheet_name]
            data = []
            for row in range(2, ws.max_row + 1):
                serial_excel = ws["A" + str(row)].value  # 序号
                name_excel = ws["B" + str(row)].value  # 网络安全设备名称
                is_virtual = ws["C" + str(row)].value  # 是否虚拟设备
                if is_virtual == 'O':
                    is_virtual = '否'
                else:
                    is_virtual = '是'
                sysversion_excel = ws["E" + str(row)].value  # 型号/系统及版本
                moder_excel = ws["G" + str(row)].value  # ip地址掩码
                if moder_excel == '':
                    moder_excel = '/'
                num_excel = ws["H" + str(row)].value  # 数量
                yongtu_excel = ws["F" + str(row)].value  # 用途
                is_hot = '否'
                if int(num_excel) > 1:
                    is_hot = '是'
                zhongyaocd_excel = str(ws["I" + str(row)].value)[:-3]  # 重要程度
                mask = '/'
                info = [serial_excel, name_excel, is_virtual, sysversion_excel, loc_excel, sushuwnagluoquyu_excel,
                        moder_excel, num_excel, yongtu_excel, is_hot, zhongyaocd_excel, mask]
                data.append(info)
            max_row = ws.max_row
            table = doc.tables[7]
            for i in range(max_row - 1):  # 模板中已经有一行了，所以总共只需增加len(supplier)行
                table.add_row()
            for row in range(1, max_row):
                for col in range(12):
                    table.cell(row, col).text = str(data[row - 1][col])

    wb.close()
    print('表8：安全设备处理【完成】')


def sheet08_docx():
    """
表10服务器设备的内容
    :param doc: 信息调查表模板
    :param location_excel: 机房物理位置
    :return fuwuqinum:服务器数量
    """
    # wb = load_workbook(path_docx)
    global fuwuqinum
    for sheet_name in wb.sheetnames:
        if '服务器-存储设备' in sheet_name:
            ws = wb[sheet_name]
            data = []
            for row in range(2, ws.max_row + 1):
                serial_excel = ws["A" + str(row)].value  # 序号
                name_excel = ws["B" + str(row)].value  # 服务器设备名称
                is_virtual = ws["C" + str(row)].value  # 是否虚拟设备
                if is_virtual == 'O':
                    is_virtual = '否'
                else:
                    is_virtual = '是'
                sysversion_excel = '服务器区'  # 物理/逻辑区域
                moder_excel = ws["D" + str(row)].value  # 操作系统及版本
                app_excel = ws["H" + str(row)].value  # IP地址
                if not app_excel:
                    app_excel = '/'
                impor_excel = ws["F" + str(row)].value  # 数据库管理系统及版本
                zhongjianjian_excel = ws["G" + str(row)].value  # 中间件及版本
                num_excel = ws["I" + str(row)].value  # 数量
                is_hot = '否'  # 是否热备默认是否，当服务器数量大于1时候是热备
                if int(num_excel) > 1:
                    is_hot = '是'
                inport = str(ws["J" + str(row)].value)[:-3]  # 重要程度
                info = [serial_excel, name_excel, is_virtual, sysversion_excel, moder_excel, app_excel, impor_excel,
                        zhongjianjian_excel, appname, num_excel, is_hot, inport]
                data.append(info)
            max_row = ws.max_row
            fuwuqinum = max_row - 1
            table = doc.tables[8]
            for i in range(max_row - 1):  # 模板中已经有一行了，所以总共只需增加len(supplier)行
                table.add_row()
            for row in range(1, max_row):
                for col in range(12):
                    table.cell(row, col).text = str(data[row - 1][col])

            print('服务器数量：', fuwuqinum)
    wb.close()
    print('表7：服务器情况处理【完成】')
    return fuwuqinum


def sheet09_docx():
    """
表11终端设备的内容
    :param doc: 信息调查表模板
    :param location_excel: 机房物理位置
    :return zhduangnum:终端数量
    """
    global zhduangnum
    # wb = load_workbook(path_docx)
    data = []
    for sheet_name in wb.sheetnames:
        if '终端-感知设备-现场设备' in sheet_name:
            ws = wb[sheet_name]
            for row in range(2, ws.max_row + 1):
                serial_excel = ws["A" + str(row)].value  # 序号
                name_excel = ws["B" + str(row)].value  # 终端名称
                is_virtual = ws["C" + str(row)].value  # 是否虚拟设备
                if is_virtual == 'O':
                    is_virtual = '否'
                else:
                    is_virtual = '是'
                moder1_excel = ws["D" + str(row)].value  # 操作系统/控制软件及版本
                sysversion1_excel = ws["F" + str(row)].value  # ip地址
                if sysversion1_excel == '':
                    sysversion1_excel = '/'
                num_excel = ws["G" + str(row)].value  # 数量/台
                app1_excel = ws["E" + str(row)].value  # 用途
                impor1_excel = str(ws["H" + str(row)].value)[:-3]  # 重要程度
                info = [serial_excel, name_excel, is_virtual, loc_excel, moder1_excel, sysversion1_excel, num_excel,
                        app1_excel, impor1_excel]
                data.append(info)
            max_row = ws.max_row
            zhduangnum = max_row - 1
            table = doc.tables[9]
            for i in range(max_row - 1):  # 模板中已经有一行了，所以总共只需增加len(supplier)行
                table.add_row()
            for row in range(1, max_row):
                for col in range(9):
                    table.cell(row, col).text = str(data[row - 1][col])

    wb.close()
    print('表9:终端处理【完成】')
    return zhduangnum


def sheet10db_docx():
    """
表10系统管理软件/平台情况 (数据库情况) 。
    :param doc:
    :param fuwuqinum: 服务器的数量，主要是为了确定DOC文档从那一列开始写数据
    :param loc_excel: 有一列要填写物理位置，多以要用到这个参数
    """
    # wb = load_workbook(path_docx)
    data_db = []
    db_max_row = 0
    for sheet_name in wb.sheetnames:
        if '数据库管理系统' in sheet_name:
            ws = wb[sheet_name]
            for row in range(2, ws.max_row + 1):
                # serialdb_excel = ws["A" + str(row)].value  # 数据库序号
                namedb_excel = ws["B" + str(row)].value  # 数据库名称
                moderdb_excel = ws["H" + str(row)].value  # 所在设备名称
                appdb_excel = ws["D" + str(row)].value  # 版本
                if appdb_excel == '':
                    appdb_excel = '/'
                gongnengdb_excel = '数据持久化存储'  # 主要功能
                impordb_excel = str(ws["J" + str(row)].value)[:-3]  # 重要程度
                infodb = [namedb_excel, moderdb_excel, appdb_excel, gongnengdb_excel, impordb_excel]
                data_db.append(infodb)
            db_max_row = ws.max_row
    wb.close()

    return db_max_row, data_db


def sheet10soft_docx():
    """
表10系统管理软件/平台情况 (系统管理软件情况) 。
    :param doc:
    :param fuwuqinum: 服务器的数量，主要是为了确定DOC文档从那一列开始写数据
    :param loc_excel: 有一列要填写物理位置，多以要用到这个参数
    """
    # wb = load_workbook(path_docx)
    data_soft = []
    soft_max_row = 0
    for sheet_name in wb.sheetnames:
        if '系统管理平台-全局扩展' in sheet_name:
            ws = wb[sheet_name]
            data = []
            for row in range(2, ws.max_row + 1):
                # serial_excel = ws["A" + str(row)].value  # 序号
                name_excel = ws["B" + str(row)].value  # 名称
                moder_excel = ws["C" + str(row)].value  # 所在设备名称
                app_excel = ws["D" + str(row)].value  # 版本
                if app_excel == None:
                    app_excel = '/'
                gongneng_excel = ws["E" + str(row)].value  # 主要功能
                impor_excel = str(ws["G" + str(row)].value)[:-3]  # 重要程度
                info = [name_excel, moder_excel, app_excel, gongneng_excel, impor_excel]
                data_soft.append(info)
            soft_max_row = ws.max_row

    wb.close()

    return soft_max_row, data_soft


def sheet10_docx():
    db_max_row, data_db = sheet10db_docx()
    soft_max_row, data_soft = sheet10soft_docx()
    xuhao = int(db_max_row) + int(soft_max_row)
    for i in data_db:
        data_soft.append(i)

    max_row = xuhao
    table = doc.tables[10]
    for i in range(max_row - 2):
        table.add_row()
    for row in range(1, max_row - 1):
        for col in range(5):
            table.cell(row, col + 1).text = str(data_soft[row - 1][col])
    for row in range(1, max_row - 1):  # 添加序号
        table.cell(row, 0).text = str(row)
    print('表11：系统管理软件/平台情况 处理【完成】 ')


def sheet11_docx(sys_name):
    """
表13：业务应用软件
    :param doc: 源信息调查表
    :return: 应用系统软件名称
    """
    global appname
    for sheet_name in wb.sheetnames:
        if '业务应用软件-平台' in sheet_name:
            ws = wb[sheet_name]
            data = []
            for row in range(2, ws.max_row + 1):
                serial_excel = ws["A" + str(row)].value  # 业务应用软件序号
                APP_excel = ws["B" + str(row)].value  # 业务应用系统/平台名称
                yewu_excel = ws["D" + str(row)].value  # 业务应用软件及版本
                yinorsoft_excel = '软件'  # 硬件/软件平台
                moshi_excel = 'B/S'  # 应用模式（C/S或B/S）
                gongneng_excel = ws["C" + str(row)].value  # 主要功能
                is_waibao_excel = '外包软件开发/' + ws["E" + str(row)].value  # 自行开发/外包开发及开发商
                yongtu_excel = '内部和外部用户/100+'  # 主要用户类别及数量
                is_import = str(ws["F" + str(row)].value)[:-3]  # 重要程度

                info = [serial_excel, APP_excel, sys_name, yewu_excel, yinorsoft_excel,
                        moshi_excel, gongneng_excel, is_waibao_excel, yongtu_excel, is_import]

                data.append(info)
            appname = data[0][1]
            max_row = ws.max_row
            table = doc.tables[11]
            for i in range(max_row - 1):
                table.add_row()
            for row in range(1, max_row):
                for col in range(10):
                    table.cell(row, col).text = str(data[row - 1][col])

    wb.close()
    print('表13:业务应用软件处理【完成】')
    print("appname", appname)

    return appname


def sheet13_docx():
    """
表13：数据资源
    :param doc: 源信息调查表
    :return: 应用系统软件名称
    """
    data = []
    for sheet_name in wb.sheetnames:
        if '关键数据类别' in sheet_name:
            ws = wb[sheet_name]
            for row in range(2, ws.max_row + 1):
                bwk = {'保密性': '×', '完整性': '×', '可用性': '√'}
                serial_excel = ws["A" + str(row)].value  # 序号
                APP_excel = ws["B" + str(row)].value  # 数据类别
                yewu_excel = ws["C" + str(row)].value  # 所属业务应用
                if yewu_excel == None or yewu_excel == '':
                    yewu_excel = '/'
                gongneng_excel = ws["D" + str(row)].value  # 数据安全性要求
                if gongneng_excel:
                    temp = gongneng_excel.split('、')
                    for i in temp:
                        if i in bwk:
                            bwk[i] = '√'

                is_waibao_excel = '定时任务自动备份'  # 备份方式
                yongtu_excel = '每天全量备份'  # 备份频率
                beifenjiezhi_excel = '移动硬盘'  # 备份介质
                is_import = str(ws["L" + str(row)].value)[:-3]  # 重要程度
                info = [serial_excel, APP_excel, yewu_excel, bwk['保密性'], bwk['完整性'], bwk['可用性'],
                        is_waibao_excel, yongtu_excel, beifenjiezhi_excel, is_import]

                data.append(info)

            max_row = ws.max_row
            table = doc.tables[13]
            for i in range(max_row - 1):
                table.add_row()
            for row in range(1, max_row):
                for col in range(10):
                    table.cell(row + 1, col).text = str(data[row - 1][col])
    wb.close()
    print('表13:数据资源处理【完成】')

    return appname


def sheet15_docx():
    '''
表格16 安全管理文档情况
    '''
    data = []
    for sheet_name in wb.sheetnames:
        if '安全管理文档' in sheet_name:
            ws = wb[sheet_name]
            for row in range(2, ws.max_row + 1):
                serial_excel = ws["A" + str(row)].value  # 序号
                APP_excel = ws["B" + str(row)].value  # 文档名称
                yewu_excel = ws["C" + str(row)].value  # 主要内容
                mask_excel = '无'  # 备注
                info = [serial_excel, APP_excel, yewu_excel, mask_excel]
                data.append(info)
            max_row = ws.max_row
            table = doc.tables[15]
            for i in range(max_row - 1):
                table.add_row()
            for row in range(1, max_row):
                for col in range(4):
                    table.cell(row, col).text = str(data[row - 1][col])

    wb.close()
    wb.close()
    print('表13:安全管理文档情况处理【完成】')


def add_pic(asset_pic):
    '''

    :param doc:
    :param asset_pic: 图片地址
    :return:
    '''
    if asset_pic:
        doc.add_picture(asset_pic, width=shared.Cm(15.34))
        print('拓扑图插入【成】')
    else:
        print('无图片,拓扑图插入【未完成】')


def add_tpt_desc():
    '''
    添加网络拓扑图描述
    :return:
    '''


def start_up(path_docx, path_doc, bumen, sys_name, outputpath):
    '''
    :param path_docx: 用户上传的资产表excel路径
    :param path_doc: 信息调查表模板路径
    :return:
    '''
    global doc, wb
    wb = load_workbook(path_docx)
    doc = Document(path_doc)
    sheet11_docx(sys_name)  # 业务应用软件
    sheet02_docx()  # 物理环境
    sheet08_docx()  # 服务器
    sheet09_docx()  # 终端
    sheet01_docx()
    sheet04_docx(bumen)
    sheet05_docx()
    sheet06_docx()
    sheet07_docx()
    sheet10_docx()
    sheet13_docx()
    sheet15_docx()
    # add_pic(picture_path)
    doc.save(outputpath)
    return outputpath

# if __name__ == '__main__':
#     start_up()
