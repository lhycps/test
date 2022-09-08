import re
from docx import Document

docxpath = './demo.docx'
doc = Document(docxpath)

restr = '本次等级测评分为四个过程：测评准备过程、方案编制过程、测评实施过程、分析与报告编制过程。*'
pos = 0

for i in range(len(doc.paragraphs)):
    matchRet = re.findall(restr, doc.paragraphs[i].text)
    if len(matchRet) > 0:
        pos = i

print(doc.paragraphs[pos + 1].text)
print(doc.paragraphs[pos + 2].text)
print(doc.paragraphs[pos + 3].text)
print(doc.paragraphs[pos + 4].text)