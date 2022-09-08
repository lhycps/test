# encoding: GBK
import re

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from copy import deepcopy
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.ns import qn


def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)
    return tbl, p


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


# ���ñ��ı߿�
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def inser_page(doc):
    paragraphs = list(doc.paragraphs)
    for i in range(len(paragraphs)):
        if i == 15:
            par1 = doc.paragraphs[i].insert_paragraph_before('���ⵥλ����ȷ��ǩ��ҳ\n')
            par1.style = doc.styles['Title']


def insert_table_after_text(document, table):
    '''
    �����������ǩ������
    @param document:
    @param table:
    @return:
    '''
    inser_page(document)
    # tbl, p = copy_table_after(table=table, paragraph=paragraph_target)
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text
        if paragraph_text == '���ⵥλ����ȷ��ǩ��ҳ\n':
            move_table_after(table, paragraph)
            # ��������ӱ߿�
            for r, row in enumerate(table.rows):  # ����������
                for c, cell in enumerate(row.cells):  # ����ÿһ�е���
                    cell1 = (r, c)  # �ڵ�Ԫ����д�뵱ǰ��λ�ã�����С��У�
                    set_cell_border(
                        table.cell(r, c),
                        top={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                        bottom={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                        left={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                        right={"sz": 1, "val": "single", "color": "#000000", "space": "0"},
                    )


def del_parcontent(doc):
    """
ɾ��ָ����������
    :param doc:
    """


    restr_start = '�ȼ�����������������ͼ'
    pos = 0
    for i in range(len(doc.paragraphs)):
        matchRet = re.findall(restr_start, doc.paragraphs[i].text)
        if len(matchRet) > 0:
            pos = i
    runspaelen = len(doc.paragraphs[pos - 2].runs)
    delstartpos = 0
    for i in range(runspaelen - 1):
        if doc.paragraphs[pos - 2].runs[i].text == '���У�':
            delstartpos = i
    for j in range(delstartpos, runspaelen - 1):
        doc.paragraphs[pos - 2].runs[j].clear()

    # doc.paragraphs[pos - 2].clear()  # ɾ��ָ������


# if __name__ == '__main__':
#     path_word = '../01��̬�ļ�_��Ҫ����/���������û�����/LISS2021DJCP081_qqqq��_20210629150407544����ࣩ.docx'
#     path_word01 = '../02��̬�ļ�_CNAS/��������+��������/����ǩ��ҳ.docx'
#     doc = Document(path_word)
#     doc_target = Document(path_word01)
#     paragraph_target = doc_target.paragraphs[0]
#     table = doc_target.tables[0]
#     insert_table_after_text(doc)
#     doc.save('./5555.docx')
